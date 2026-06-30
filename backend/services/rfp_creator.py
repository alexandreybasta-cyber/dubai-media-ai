"""
RFP Creator Service - AI-powered RFP document generation using Qwen via DashScope.
"""

import io
import json
import logging
from datetime import datetime

import httpx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from config import settings

logger = logging.getLogger(__name__)

RFP_SECTIONS = [
    "Executive Summary",
    "Organization Background",
    "Scope of Work",
    "Technical Requirements & Specifications",
    "Evaluation Criteria Matrix",
    "Timeline & Milestones",
    "Budget & Commercial Terms",
    "Compliance & Regulatory Requirements",
    "Submission Guidelines & Instructions",
    "Terms & Conditions",
]

SYSTEM_PROMPT = """You are a professional RFP (Request for Proposal) document writer for Dubai Media Incorporated, 
a leading media organization in the UAE. You produce clear, formal, comprehensive RFP documents that follow 
international procurement standards while adhering to UAE regulations and business practices.

Your writing style:
- Professional and precise language
- Clear structure with actionable requirements
- Industry-standard terminology
- Compliant with UAE procurement regulations
- Suitable for international vendors

When generating content, produce detailed, substantive paragraphs — not bullet points unless specifically requested.
Each section should be 150-300 words minimum to ensure completeness."""

BILINGUAL_INSTRUCTION = """
After generating the English content, provide an Arabic translation of the same content.
Format: First the English text, then a separator "---AR---", then the Arabic translation.
The Arabic should be professional and suitable for official UAE government documents."""


class RFPCreator:
    """AI-powered RFP document generator using Qwen via DashScope."""

    def __init__(self):
        self.api_url = f"{settings.DASHSCOPE_BASE_URL}/chat/completions"
        self.model = settings.MODEL_TEXT
        self.api_key = settings.DASHSCOPE_API_KEY
        self.max_retries = 3

    async def _call_llm(self, messages: list[dict], temperature: float = 0.7) -> str:
        """Call DashScope API with retries and exponential backoff."""
        if not self.api_key:
            raise ValueError(
                "DASHSCOPE_API_KEY is not configured. Please set it in your .env file."
            )

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": self.model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": 4096,
        }

        last_error = None
        for attempt in range(self.max_retries):
            try:
                async with httpx.AsyncClient(timeout=120.0) as client:
                    response = await client.post(
                        self.api_url, headers=headers, json=payload
                    )
                    if response.status_code == 200:
                        data = response.json()
                        return data["choices"][0]["message"]["content"]
                    else:
                        last_error = f"API returned status {response.status_code}: {response.text}"
                        logger.warning(
                            f"DashScope API attempt {attempt + 1} failed: {last_error}"
                        )
            except httpx.TimeoutException as e:
                last_error = f"Request timed out: {str(e)}"
                logger.warning(f"DashScope API attempt {attempt + 1} timed out")
            except Exception as e:
                last_error = str(e)
                logger.warning(
                    f"DashScope API attempt {attempt + 1} error: {last_error}"
                )

            if attempt < self.max_retries - 1:
                import asyncio
                await asyncio.sleep(2 ** attempt)

        raise RuntimeError(f"DashScope API failed after {self.max_retries} attempts: {last_error}")

    async def generate_rfp(self, input_data: dict) -> dict:
        """Generate complete RFP from structured input."""
        sections = []
        bilingual = input_data.get("language") == "both"
        tone_instruction = f"Use a {input_data.get('tone', 'formal')} tone throughout."

        for section_name in RFP_SECTIONS:
            content = await self._generate_section(
                section_name, input_data, tone_instruction, bilingual
            )
            section_entry = {"name": section_name, "content_en": content}
            if bilingual and "---AR---" in content:
                parts = content.split("---AR---")
                section_entry["content_en"] = parts[0].strip()
                section_entry["content_ar"] = parts[1].strip()
            elif input_data.get("language") == "ar":
                section_entry["content_ar"] = content
                section_entry["content_en"] = ""
            sections.append(section_entry)

        return {
            "project_title": input_data.get("project_title", ""),
            "language": input_data.get("language", "en"),
            "tone": input_data.get("tone", "formal"),
            "sections": sections,
            "generated_at": datetime.utcnow().isoformat(),
            "input_data": input_data,
        }

    async def _generate_section(
        self,
        section_name: str,
        input_data: dict,
        tone_instruction: str,
        bilingual: bool,
    ) -> str:
        """Generate a single RFP section."""
        context = self._build_section_context(section_name, input_data)
        lang_instruction = BILINGUAL_INSTRUCTION if bilingual else ""
        if input_data.get("language") == "ar":
            lang_instruction = "Generate the content in Arabic suitable for official UAE government documents."

        user_prompt = f"""Generate the "{section_name}" section for an RFP document.

Project Title: {input_data.get('project_title', 'N/A')}
Project Overview: {input_data.get('project_overview', 'N/A')}

{context}

{tone_instruction}
{lang_instruction}

Generate professional, detailed content for this section. Do not include the section title itself — just the body content."""

        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ]
        return await self._call_llm(messages)

    def _build_section_context(self, section_name: str, input_data: dict) -> str:
        """Build section-specific context from input data."""
        context_parts = []

        if section_name == "Scope of Work":
            if input_data.get("scope_of_work"):
                context_parts.append(
                    f"Scope details provided:\n{input_data['scope_of_work']}"
                )

        elif section_name == "Technical Requirements & Specifications":
            reqs = input_data.get("technical_requirements", [])
            if reqs:
                context_parts.append(
                    "Technical requirements:\n"
                    + "\n".join(f"- {r}" for r in reqs)
                )

        elif section_name == "Evaluation Criteria Matrix":
            criteria = input_data.get("evaluation_criteria", [])
            if criteria:
                context_parts.append("Evaluation criteria:")
                for c in criteria:
                    context_parts.append(
                        f"- {c.get('name', 'N/A')} (Weight: {c.get('weight', 0)}%): {c.get('description', '')}"
                    )
                context_parts.append(
                    "Present these as a weighted evaluation matrix table."
                )

        elif section_name == "Timeline & Milestones":
            timeline = input_data.get("timeline", {})
            if timeline:
                context_parts.append(f"Start date: {timeline.get('start_date', 'TBD')}")
                context_parts.append(f"End date: {timeline.get('end_date', 'TBD')}")
                milestones = timeline.get("milestones", [])
                if milestones:
                    context_parts.append("Milestones:")
                    for m in milestones:
                        context_parts.append(
                            f"- {m.get('name', 'N/A')}: {m.get('date', 'TBD')}"
                        )

        elif section_name == "Budget & Commercial Terms":
            budget = input_data.get("budget_range")
            if budget:
                context_parts.append(
                    f"Budget range: {budget.get('currency', 'AED')} "
                    f"{budget.get('min', 'N/A')} - {budget.get('max', 'N/A')}"
                )
            else:
                context_parts.append(
                    "No specific budget range provided. Include standard commercial terms."
                )

        elif section_name == "Compliance & Regulatory Requirements":
            compliance = input_data.get("compliance_requirements", [])
            if compliance:
                context_parts.append(
                    "Compliance requirements:\n"
                    + "\n".join(f"- {r}" for r in compliance)
                )

        elif section_name == "Organization Background":
            industry = input_data.get("industry", "Media")
            context_parts.append(
                f"Industry context: {industry}. "
                "Dubai Media Incorporated is a government media organization under the "
                "Government of Dubai, responsible for managing and operating Dubai's media sector."
            )

        return "\n".join(context_parts)

    async def regenerate_section(
        self, rfp_data: dict, section_name: str, instructions: str = ""
    ) -> str:
        """Regenerate a single section of the RFP."""
        input_data = rfp_data.get("input_data", {})
        tone = input_data.get("tone", "formal")
        language = input_data.get("language", "en")
        bilingual = language == "both"

        existing_content = ""
        for section in rfp_data.get("sections", []):
            if section["name"] == section_name:
                existing_content = section.get("content_en", "")
                break

        lang_instruction = BILINGUAL_INSTRUCTION if bilingual else ""
        if language == "ar":
            lang_instruction = "Generate the content in Arabic suitable for official UAE government documents."

        user_prompt = f"""Regenerate the "{section_name}" section for an RFP document.

Project Title: {rfp_data.get('project_title', 'N/A')}

Previous content for reference:
{existing_content}

Additional instructions: {instructions or 'Improve and expand the content while maintaining professional quality.'}

Use a {tone} tone.
{lang_instruction}

Generate professional, detailed content for this section. Do not include the section title — just the body content."""

        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ]
        content = await self._call_llm(messages)
        return content

    def export_docx(self, rfp_data: dict) -> bytes:
        """Generate DOCX file from RFP data."""
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("DUBAI MEDIA INCORPORATED")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)

        doc.add_paragraph()  # spacer

        rfp_title = doc.add_paragraph()
        rfp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rfp_run = rfp_title.add_run("REQUEST FOR PROPOSAL")
        rfp_run.bold = True
        rfp_run.font.size = Pt(20)

        doc.add_paragraph()

        project_title = doc.add_paragraph()
        project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_run = project_title.add_run(rfp_data.get("project_title", "Untitled RFP"))
        pt_run.bold = True
        pt_run.font.size = Pt(14)

        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")

        doc.add_page_break()

        # Table of Contents
        toc_heading = doc.add_heading("Table of Contents", level=1)
        for i, section in enumerate(rfp_data.get("sections", []), 1):
            toc_para = doc.add_paragraph(f"{i}. {section['name']}", style="List Number")

        doc.add_page_break()

        # Sections
        for i, section in enumerate(rfp_data.get("sections", []), 1):
            doc.add_heading(f"{i}. {section['name']}", level=1)

            content = section.get("content_en", "") or section.get("content_ar", "")

            if section["name"] == "Evaluation Criteria Matrix":
                self._add_criteria_table_docx(doc, rfp_data, content)
            elif section["name"] == "Timeline & Milestones":
                self._add_timeline_table_docx(doc, rfp_data, content)
            else:
                for paragraph_text in content.split("\n\n"):
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text.strip())

            # Add Arabic content if bilingual
            if section.get("content_ar"):
                doc.add_paragraph()
                ar_heading = doc.add_paragraph()
                ar_run = ar_heading.add_run("النسخة العربية")
                ar_run.bold = True
                ar_para = doc.add_paragraph(section["content_ar"])
                ar_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph()  # spacing

        # Footer with page numbers
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = "Dubai Media Incorporated - Confidential"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_criteria_table_docx(self, doc, rfp_data: dict, content: str):
        """Add evaluation criteria as a formatted table."""
        input_data = rfp_data.get("input_data", {})
        criteria = input_data.get("evaluation_criteria", [])

        if criteria:
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "#"
            hdr_cells[1].text = "Criterion"
            hdr_cells[2].text = "Weight (%)"
            hdr_cells[3].text = "Description"

            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for i, c in enumerate(criteria, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = c.get("name", "")
                row_cells[2].text = str(c.get("weight", 0))
                row_cells[3].text = c.get("description", "")

            doc.add_paragraph()

        # Also add the generated content
        if content:
            for paragraph_text in content.split("\n\n"):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

    def _add_timeline_table_docx(self, doc, rfp_data: dict, content: str):
        """Add timeline as a table."""
        input_data = rfp_data.get("input_data", {})
        timeline = input_data.get("timeline", {})
        milestones = timeline.get("milestones", [])

        if milestones:
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "#"
            hdr_cells[1].text = "Milestone"
            hdr_cells[2].text = "Target Date"

            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for i, m in enumerate(milestones, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = m.get("name", "")
                row_cells[2].text = m.get("date", "TBD")

            doc.add_paragraph()

        if content:
            for paragraph_text in content.split("\n\n"):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

    def export_pdf(self, rfp_data: dict) -> bytes:
        """Generate PDF file from RFP data."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                "RFPTitle",
                parent=styles["Title"],
                fontSize=24,
                spaceAfter=12,
                textColor=colors.HexColor("#F97316"),
            )
        )
        styles.add(
            ParagraphStyle(
                "RFPSubtitle",
                parent=styles["Title"],
                fontSize=16,
                spaceAfter=24,
            )
        )
        styles.add(
            ParagraphStyle(
                "SectionHeading",
                parent=styles["Heading1"],
                fontSize=14,
                spaceBefore=20,
                spaceAfter=10,
                textColor=colors.HexColor("#1F2937"),
            )
        )
        styles.add(
            ParagraphStyle(
                "RFPBody",
                parent=styles["Normal"],
                fontSize=10,
                spaceAfter=8,
                leading=14,
            )
        )

        elements = []

        # Title page
        elements.append(Spacer(1, 2 * inch))
        elements.append(Paragraph("DUBAI MEDIA INCORPORATED", styles["RFPTitle"]))
        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph("REQUEST FOR PROPOSAL", styles["RFPSubtitle"]))
        elements.append(
            Paragraph(
                rfp_data.get("project_title", "Untitled RFP"), styles["RFPSubtitle"]
            )
        )
        elements.append(Spacer(1, inch))
        elements.append(
            Paragraph(
                f"Date: {datetime.now().strftime('%B %d, %Y')}", styles["Normal"]
            )
        )
        elements.append(PageBreak())

        # Table of Contents
        elements.append(Paragraph("Table of Contents", styles["SectionHeading"]))
        for i, section in enumerate(rfp_data.get("sections", []), 1):
            elements.append(
                Paragraph(f"{i}. {section['name']}", styles["RFPBody"])
            )
        elements.append(PageBreak())

        # Sections
        for i, section in enumerate(rfp_data.get("sections", []), 1):
            elements.append(
                Paragraph(f"{i}. {section['name']}", styles["SectionHeading"])
            )

            content = section.get("content_en", "") or section.get("content_ar", "")

            if section["name"] == "Evaluation Criteria Matrix":
                self._add_criteria_table_pdf(elements, rfp_data, content, styles)
            elif section["name"] == "Timeline & Milestones":
                self._add_timeline_table_pdf(elements, rfp_data, content, styles)
            else:
                for paragraph_text in content.split("\n\n"):
                    if paragraph_text.strip():
                        # Escape XML special chars for reportlab
                        safe_text = (
                            paragraph_text.strip()
                            .replace("&", "&amp;")
                            .replace("<", "&lt;")
                            .replace(">", "&gt;")
                        )
                        elements.append(Paragraph(safe_text, styles["RFPBody"]))

            elements.append(Spacer(1, 0.3 * inch))

        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_criteria_table_pdf(self, elements, rfp_data, content, styles):
        """Add evaluation criteria table to PDF."""
        input_data = rfp_data.get("input_data", {})
        criteria = input_data.get("evaluation_criteria", [])

        if criteria:
            table_data = [["#", "Criterion", "Weight (%)", "Description"]]
            for i, c in enumerate(criteria, 1):
                table_data.append(
                    [str(i), c.get("name", ""), str(c.get("weight", 0)), c.get("description", "")]
                )

            table = Table(table_data, colWidths=[0.5 * inch, 2 * inch, 1 * inch, 3 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F97316")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                        ("TOPPADDING", (0, 0), (-1, 0), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FFF7ED")]),
                    ]
                )
            )
            elements.append(table)
            elements.append(Spacer(1, 0.2 * inch))

        if content:
            for paragraph_text in content.split("\n\n"):
                if paragraph_text.strip():
                    safe_text = (
                        paragraph_text.strip()
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    elements.append(Paragraph(safe_text, styles["RFPBody"]))

    def _add_timeline_table_pdf(self, elements, rfp_data, content, styles):
        """Add timeline table to PDF."""
        input_data = rfp_data.get("input_data", {})
        timeline = input_data.get("timeline", {})
        milestones = timeline.get("milestones", [])

        if milestones:
            table_data = [["#", "Milestone", "Target Date"]]
            for i, m in enumerate(milestones, 1):
                table_data.append([str(i), m.get("name", ""), m.get("date", "TBD")])

            table = Table(table_data, colWidths=[0.5 * inch, 4 * inch, 2 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F97316")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                        ("TOPPADDING", (0, 0), (-1, 0), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FFF7ED")]),
                    ]
                )
            )
            elements.append(table)
            elements.append(Spacer(1, 0.2 * inch))

        if content:
            for paragraph_text in content.split("\n\n"):
                if paragraph_text.strip():
                    safe_text = (
                        paragraph_text.strip()
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    elements.append(Paragraph(safe_text, styles["RFPBody"]))
